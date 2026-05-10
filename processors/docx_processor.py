from typing import Dict, List
from .base import BaseProcessor

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxProcessor(BaseProcessor):
    """Word文档处理器"""

    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx']

    def extract_text(self, file_path: str) -> str:
        """从Word文档提取文本"""
        if not DOCX_AVAILABLE:
            raise ImportError("请安装python-docx: pip install python-docx")

        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return '\n'.join(text_parts)
        except Exception as e:
            raise RuntimeError(f"DOCX解析失败: {str(e)}")

    def extract_metadata(self, file_path: str) -> Dict:
        """提取Word文档元数据"""
        info = self.get_file_info(file_path)
        metadata = {
            **info,
            'type': 'word',
            'paragraph_count': 0,
            'table_count': 0,
            'author': ''
        }

        if not DOCX_AVAILABLE:
            return metadata

        try:
            doc = Document(file_path)
            metadata['paragraph_count'] = len([p for p in doc.paragraphs if p.text.strip()])
            metadata['table_count'] = len(doc.tables)

            core_props = doc.core_properties
            metadata['author'] = getattr(core_props, 'author', '')
        except Exception:
            pass

        return metadata
